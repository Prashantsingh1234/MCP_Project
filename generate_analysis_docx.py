"""Generate MCPDischarge project analysis as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)   # headings
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)   # sub-headings
ACCENT     = RGBColor(0x00, 0x70, 0xC0)   # table headers
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
CODE_BG    = RGBColor(0xF5, 0xF5, 0xF5)
CODE_FG    = RGBColor(0x1E, 0x1E, 0x1E)
GREEN      = RGBColor(0x37, 0x86, 0x35)
ORANGE     = RGBColor(0xD6, 0x79, 0x10)


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(table):
    """Thin borders on every cell."""
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = DARK_BLUE
    p.runs[0].font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = MID_BLUE
    p.runs[0].font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = DARK_BLUE
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size  = Pt(10.5)
    r.font.bold  = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = CODE_FG
    # background shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F5F5F5")
    pPr.append(shd)
    return p


def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_border(t)
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, ACCENT)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9.5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # data rows
    for ri, row in enumerate(rows):
        bg = LIGHT_GREY if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
        t.rows[ri + 1].cells[0].paragraphs[0].runs[0].font.bold = True
    # column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return t


def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("MCPDischarge")
tr.font.size = Pt(32)
tr.font.bold = True
tr.font.color.rgb = DARK_BLUE

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr2 = tp2.add_run("Complete Project Analysis")
tr2.font.size = Pt(20)
tr2.font.color.rgb = MID_BLUE

doc.add_paragraph()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr3 = tp3.add_run(
    "Architecture · Workflow · Tools · Validation · Guardrails\n"
    "Monitoring · Evaluation · Gateway · Models · RBAC\n"
    "PHI Guard · Response Formatter · Metrics"
)
tr3.font.size  = Pt(12)
tr3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
tr3.font.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Project Overview")
body(
    "MCPDischarge is a production-grade, multi-agent healthcare automation system built on the "
    "Model Context Protocol (MCP). It automates hospital discharge workflows that traditionally "
    "required 45+ minutes and 15 manual handoffs — EHR lookups, pharmacy stock checks, billing "
    "invoice generation — and reduces them to <1 second of orchestrated tool calls."
)
body(
    "The system coordinates three hospital departments (EHR, Pharmacy, Billing) via standardised "
    "MCP tool calls while enforcing RBAC (Role-Based Access Control), PHI (Protected Health "
    "Information) boundaries, and comprehensive telemetry."
)

make_table(
    ["Metric", "Value"],
    [
        ["MCP Tools (total)", "59"],
        ["Python source files", "25+"],
        ["Lines of code", "~3,500"],
        ["Custom exceptions", "7"],
        ["RBAC roles", "4"],
        ["RBAC permission types", "30+"],
        ["PHI fields guarded", "5"],
        ["Patients (test data)", "6"],
        ["Drugs (inventory)", "17 (4 out-of-stock)"],
        ["Guardrail layers", "7"],
        ["Evaluation test files", "6"],
        ["MCP servers", "5 (ports 8001–8005)"],
        ["Gateway port", "8000"],
    ],
    col_widths=[2.5, 3.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. DIRECTORY STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Directory Structure (Every File)")

structure = """\
project5_mcp/
├── .env                              Azure OpenAI + LangSmith secrets
├── .env.example                      Safe-to-commit template
├── requirements.txt                  20+ Python deps
├── README.md                         Architecture diagrams + usage guide
│
├── src/
│   ├── __init__.py                   Public API exports
│   │
│   ├── servers/                      5 FastMCP HTTP SSE servers
│   │   ├── mcp_servers.py            Server factory + CLI launcher
│   │   ├── ehr_server.py             Patient data, PHI stripping
│   │   ├── pharmacy_server.py        Stock, pricing, dispensing, aliases
│   │   ├── billing_server.py         Charges, insurance, invoice
│   │   ├── security_server.py        RBAC checks, violation log
│   │   └── telemetry_server.py       Metrics, traces, health
│   │
│   ├── agents/
│   │   └── discharge_agent.py        Async multi-server orchestrator
│   │
│   ├── gateway/
│   │   ├── chat_gateway.py           FastAPI HTTP gateway (port 8000)
│   │   ├── llm_azure.py              Azure OpenAI SDK wrapper
│   │   ├── invoice_pdf.py            ReportLab A4 invoice PDF + HTML
│   │   └── prescription_pdf.py       Prescription PDF + HTML
│   │
│   ├── chatbot/
│   │   ├── llm_controller.py         Master dispatcher + defense layers
│   │   ├── llm_agent.py              LLM-driven tool selection
│   │   ├── llm_provider.py           Azure OpenAI chat_with_tools()
│   │   ├── mcp_client.py             Unified RBAC+retry wrapper
│   │   ├── intent_classifier.py      Intent detection + injection defense
│   │   ├── validator.py              Entity extraction (IDs, drugs, doses)
│   │   ├── preprocessor.py           Text sanitization + LLM rewrite
│   │   ├── phi_guard.py              PHI detection, stripping, denial
│   │   ├── rbac_guard.py             Per-call RBAC enforcement
│   │   ├── response_formatter.py     Human-friendly PHI-safe output
│   │   ├── workflow_engine.py        Stateful discharge orchestration
│   │   ├── conversation_manager.py   Multi-turn state (non-PHI, cap 40)
│   │   ├── metrics.py                Per-request telemetry aggregator
│   │   ├── controller.py             Legacy controller
│   │   └── cli.py                    Interactive CLI REPL
│   │
│   └── utils/
│       ├── exceptions.py             7 custom exception types
│       ├── rbac.py                   RBACEngine singleton
│       ├── telemetry.py              Telemetry singleton
│       ├── data_loader.py            DataLoader singleton (7 JSON files)
│       └── langsmith_tracing.py      LangSmith + PHI masking (411 lines)
│
├── data/
│   ├── generate_dataset.py           Script to create JSON test data
│   ├── ehr_patients.json             6 patient records (9.8 KB)
│   ├── pharmacy_inventory.json       17 drugs + aliases (13 KB)
│   ├── billing_rate_cards.json       Ward/lab charges (2.1 KB)
│   ├── insurance_contracts.json      2 insurance companies (922 B)
│   ├── patient_insurance_map.json    Patient → insurer (689 B)
│   ├── icd10_billing_codes.json      ICD-10 → DRG codes (984 B)
│   └── rbac_policies.json            RBAC matrix (2.1 KB)
│
├── evaluation/
│   ├── eval_dashboard.py             All-patient metrics run
│   ├── test_chatbot_flows.py         Multi-turn conversation tests
│   ├── test_llm_agent_guards.py      PHI + RBAC defense tests
│   ├── test_phi_guard_invoice_payloads.py  PHI field blocking tests
│   ├── test_llm_agent_tool_name_normalization.py
│   └── test_patient_ordinal_resolution.py
│
├── demo/
│   ├── demo.py                       4 discharge + 2 failure scenarios
│   ├── generate_invoice_pdf.py       Single invoice example
│   └── multi_turn_medication_flow.py Multi-turn chatbot example
│
├── frontend/                         React + Vite + TypeScript UI
│   ├── package.json
│   ├── tsconfig.json
│   └── src/ (App.tsx, components/, pages/)
│
└── configs/
    ├── fastmcp_deployment.md         HTTP SSE setup guide
    ├── azure_foundry_mcp.md          Azure AI Foundry guide
    └── rbac_design.md                RBAC architecture patterns"""

code_block(structure)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("3. System Architecture — Layer by Layer")

layers = [
    ("Browser / Mobile UI", "React + Vite + TypeScript frontend. Communicates over HTTP to the FastAPI gateway."),
    ("FastAPI Chat Gateway (port 8000)", "Main HTTP entry point. Routes: /chat, /sessions, /invoice, /prescription, /health. Manages conversation sessions (max 50, in-memory OrderedDict). CORS enabled."),
    ("LLMChatController", "Master dispatcher. Runs 7 defense layers before touching any LLM. Delegates to LLMToolCallingAgent (flexible) or WorkflowEngine (rule-based) depending on intent."),
    ("LLMToolCallingAgent", "Sends tool schemas to Azure OpenAI. LLM selects which MCP tools to call and in what order."),
    ("WorkflowEngine", "Stateful rule-based discharge orchestration. Fallback when LLM is unavailable. LangSmith-traced."),
    ("AsyncMCPToolClient", "Unified async client for all 5 MCP servers. Enforces RBAC before every call, per-server asyncio.Lock() for SSE safety, exponential-backoff retry (max 3), telemetry recording."),
    ("MCP Servers (5×)", "FastMCP HTTP SSE servers: EHR (:8001), Pharmacy (:8002), Billing (:8003), Security (:8004), Telemetry (:8005). Each tool wrapped with MCPCallTimer + LangSmith tracing."),
    ("DataLoader Singleton", "Loads all 7 JSON files at startup. Single source of truth for patients, inventory, charges, insurance, RBAC policies."),
    ("RBACEngine Singleton", "Data-driven policy enforcement from rbac_policies.json. Maps tool names → permissions. Logs violations to Telemetry."),
    ("Telemetry Singleton", "Collects MCPCall, Alert, ChatTrace, RBACViolation records. PHI scrubbed before any write. Powers Telemetry Server queries."),
    ("LangSmith Tracing", "Optional, safe-by-default. Custom input/output processors sanitize all data. Text masking preserves patient IDs, masks names."),
]

for title, desc in layers:
    h3(title)
    body(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. REQUEST LIFECYCLE / WORKFLOW
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Request Lifecycle & Workflow")

h2("Phase 1 — Input Defense (runs before any LLM call)")
make_table(
    ["Step", "Component", "What Happens"],
    [
        ["1", "intent_classifier.is_prompt_injection()", "Regex scan: 'ignore rbac', 'bypass rbac', 'jailbreak', 'show all patients', 'system prompt'"],
        ["2", "phi_guard.deny_if_phi_requested()", "Blocks requests for patient name, DOB, MRN, discharge note, attending physician"],
        ["3", "preprocessor.sanitize_user_text()", "Unicode normalisation, zero-width char removal, patient ID normalisation (pat001→PAT-001), typo fix"],
        ["4", "preprocessor._should_llm_rewrite()", "Heuristic: 6+ non-ASCII? unbalanced quotes? material changes? → trigger LLM polish"],
        ["5", "intent_classifier.classify_intent()", "Returns intent: invalid_input | rbac_sensitive_request | bulk_request | meds_fetch | discharge_workflow | …"],
        ["6", "validator.validate_message()", "Extracts patient IDs (PAT-XXX), ordinals ('3rd patient'→PAT-003), drug names, doses; validates required fields per intent"],
    ],
    col_widths=[0.4, 2.5, 3.6],
)

h2("Phase 2 — LLM / Tool Selection")
make_table(
    ["Step", "Component", "What Happens"],
    [
        ["7", "LLMToolCallingAgent.decide_tools()", "Calls Azure OpenAI with full MCP tool schemas + user intent → returns sequence of tool calls"],
        ["8", "Fallback (_fallback())", "If LLM times out, rule-based routing: meds_fetch → get_discharge_medications directly"],
    ],
    col_widths=[0.4, 2.5, 3.6],
)

h2("Phase 3 — MCP Tool Execution")
make_table(
    ["Step", "Component", "What Happens"],
    [
        ["9",  "MCPClient.ehr_call / pharmacy_call / billing_call()", "RBAC check → acquire server lock → call MCP SSE tool → record telemetry → release lock"],
        ["10", "RBACEngine.check_permission()", "Looks up rbac_policies.json → raises RBACError if denied → logs violation"],
        ["11", "MCPCallTimer context manager", "Measures wall-clock duration, records success/failure to Telemetry"],
    ],
    col_widths=[0.4, 2.5, 3.6],
)

h2("Phase 4 — Discharge Workflow (discharge_workflow intent)")
code_block("""\
EHR: get_discharge_medications(patient_id)
  └─→ For each medication:
        Pharmacy: check_stock(drug_name, quantity, dose)
          └─→ if unavailable: Pharmacy: get_alternative(drug_name)
        Pharmacy: get_price(drug_name, quantity)
EHR: get_billing_safe_summary(patient_id)        ← PHI stripped here
Billing: validate_billing_data(billing_safe_ehr) ← validates PHI absent
Billing: get_charges(ward, los_days)
Billing: get_insurance(patient_id)
Billing: calculate_insurance_coverage(patient_id, charges)
Billing: generate_invoice(patient_id, billing_safe_ehr, drug_charges)""")

h2("Phase 5 — Response Assembly")
make_table(
    ["Step", "Component", "What Happens"],
    [
        ["12", "response_formatter.py", "Converts raw dicts → human-readable text (symbols, structured sections)"],
        ["13", "ConversationManager", "Stores non-PHI state for follow-up queries (capped 40 entries)"],
        ["14", "Telemetry.record_chat_trace()", "Logs conversation outcome (no raw user text, no PHI)"],
        ["15", "langsmith_tracing.traceable_safe()", "Sends masked trace to LangSmith (if configured)"],
        ["16", "Gateway", "Returns ControllerResponse(success, answer, data) to HTTP client"],
    ],
    col_widths=[0.4, 2.5, 3.6],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. COMPLETE TOOL LIST
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Complete Tool List (All 59 MCP Tools)")

h2("EHR Server — port 8001 (18 Tools)")
make_table(
    ["Tool", "Permission Required", "Returns"],
    [
        ["get_patient_discharge_summary",    "read_discharge_note",     "Full patient record (PHI included)"],
        ["get_discharge_medications",        "read_medications",        "Medication list with doses/routes"],
        ["get_diagnosis_codes",              "read_diagnosis_codes",    "ICD-10 codes only"],
        ["get_admission_info",               "read_admission_dates",    "Ward, dates, LOS (non-PHI)"],
        ["list_patients",                    "list_patients",           "All patient IDs"],
        ["get_billing_safe_summary",         "read_billing_safe",       "PHI-stripped patient summary"],
        ["validate_prescription",            "read_medications",        "Completeness check result"],
        ["check_drug_interactions",          "read_medications",        "Drug-drug interaction flags"],
        ["check_dose_validity",              "read_medications",        "Formulary conformance check"],
        ["update_prescription",              "update_medications",      "Updated med list confirmation"],
        ["mark_patient_ready_for_discharge", "update_discharge_status", "Status update confirmation"],
        ["get_patient_history",              "read_patient_history",    "Past medications + diagnoses"],
        ["mark_urgent_request",              "mark_urgent",             "Urgency flag confirmation"],
        ["escalate_to_doctor",               "escalate_clinical",       "Escalation confirmation"],
        ["request_represcription",           "request_represcription",  "Alternative request confirmation"],
        ["notify_patient",                   "send_notifications",      "Notification confirmation"],
        ["notify_doctor",                    "send_notifications",      "Paging confirmation"],
        ["validate_patient_id",              "read_admission_dates",    "Existence check bool"],
    ],
    col_widths=[2.4, 1.9, 2.2],
)

h2("Pharmacy Server — port 8002 (18 Tools)")
make_table(
    ["Tool", "Permission Required", "Returns"],
    [
        ["check_stock",                      "check_stock",         "found, in_stock, units, dose_conflict, alternatives"],
        ["check_bulk_stock",                 "check_stock",         "Per-drug stock dict"],
        ["list_in_stock_drugs",              "list_formulary",      "In-stock drug names"],
        ["get_alternative",                  "get_alternatives",    "Therapeutic substitute in stock"],
        ["get_all_alternatives",             "get_alternatives",    "All alternatives list"],
        ["check_therapeutic_equivalence",    "check_equivalence",   "Clinical equivalence bool"],
        ["resolve_drug_name_alias",          "resolve_alias",       "Generic name from brand/alias"],
        ["semantic_drug_search",             "search_formulary",    "Fuzzy-matched drug list"],
        ["get_price",                        "get_drug_price",      "Unit price + total in INR"],
        ["get_bulk_price",                   "get_drug_price",      "Per-drug pricing dict"],
        ["dispense_request",                 "dispense_medication", "Single dispense request record"],
        ["create_dispense_request",          "dispense_medication", "Bulk dispense request record"],
        ["confirm_dispense",                 "dispense_medication", "Confirmation status"],
        ["update_stock",                     "update_inventory",    "Updated stock confirmation"],
        ["check_nearby_pharmacy_availability","check_external",     "External availability status"],
        ["detect_dose_conflict",             "check_stock",         "Mismatch detail or clear"],
        ["flag_controlled_substance",        "flag_controlled",     "Schedule classification"],
        ["validate_drug_name",               "resolve_alias",       "Formulary membership bool"],
    ],
    col_widths=[2.4, 1.9, 2.2],
)

h2("Billing Server — port 8003 (11 Tools)")
make_table(
    ["Tool", "Permission Required", "Returns"],
    [
        ["get_charges",               "read_charge_codes",    "Ward/lab line items"],
        ["get_charges_by_icd",        "read_charge_codes",    "ICD-10→DRG charge mapping"],
        ["get_total_cost",            "read_charge_codes",    "Full bill breakdown"],
        ["get_insurance",             "read_insurance",       "Plan type, copay, deductible, max covered"],
        ["calculate_insurance_coverage","calculate_coverage", "Coverage amount + patient liability"],
        ["validate_insurance",        "validate_insurance",   "Eligibility check"],
        ["generate_payment_link",     "generate_payment_link","Payment URL"],
        ["mark_invoice_paid",         "mark_paid",            "Status update"],
        ["validate_billing_data",     "validate_billing",     "PHI absence check (raises error if PHI found)"],
        ["audit_invoice",             "audit_billing",        "Compliance audit trail"],
        ["generate_invoice",          "generate_invoice",     "Full invoice with all line items + insurance"],
    ],
    col_widths=[2.4, 1.9, 2.2],
)

h2("Security Server — port 8004 (4 Tools)")
make_table(
    ["Tool", "Returns"],
    [
        ["check_access(role, server, tool)",           "{allowed: bool, required_permission: str}"],
        ["get_role_permissions(role)",                 "{permissions_by_server, callable_tools_by_server}"],
        ["log_rbac_violation(role, tool, server, patient_id)", "Violation log entry"],
        ["get_access_logs(limit=50)",                  "Recent violation list"],
    ],
    col_widths=[3.0, 3.5],
)

h2("Telemetry Server — port 8005 (8 Tools)")
make_table(
    ["Tool", "Returns"],
    [
        ["get_mcp_call_count(patient_id?)",   "Total + by-server + success/failure counts"],
        ["get_alerts(patient_id?, level?)",   "Filtered alert list"],
        ["get_system_health()",               "TCP port ping results + uptime + aggregate metrics"],
        ["get_summary()",                     "Full telemetry aggregate"],
        ["record_chat_trace(...)",            "Confirmation (no raw user text stored)"],
        ["get_chat_traces(limit=100)",        "Recent chat interactions"],
        ["get_recent_calls(limit=100)",       "Tool calls + RBAC violations + alerts"],
        ["trace_workflow(patient_id)",        "Full ordered MCP call sequence for patient"],
    ],
    col_widths=[3.0, 3.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. VALIDATION SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Validation System")
h2("src/chatbot/validator.py — Full Pipeline")

body("validate_message(text, intent) runs the following sub-validators in sequence:", bold=True)

validators = [
    ("extract_patient_ids(text)", "Regex: PAT-\\d{3} (case-insensitive). Returns list of matched IDs."),
    ("extract_patient_ordinal(text)", "'3rd patient' → 3, 'second patient' → 2. Supports typos (frist→1, secound→2). patient_id_from_ordinal(3) → 'PAT-003'."),
    ("validate_patient_id_exists(patient_id, known_ids)", "Checks against DataLoader.patients list. Raises ValidationError if not found."),
    ("extract_drug_name(text)", "Patterns: 'check stock for X', 'is X available', 'X 5mg'. Returns best-effort drug name."),
    ("extract_dose(text)", "Regex: \\d+(\\.\\d+)?mg → '5mg'. Returns first match or None."),
    ("validate_required_patient_id(intent, patient_ids)", "Intents like discharge_workflow require exactly one patient ID. Raises ValidationError otherwise."),
    ("validate_not_bulk_for_intent(intent, patient_ids)", "Some intents (single meds_fetch) forbid bulk queries. Raises ValidationError if >1 IDs found."),
    ("resolves_to_previous_patient(text)", "Detects follow-up references ('first patient', 'same patient') for multi-turn context resolution."),
]
make_table(
    ["Function", "Behaviour"],
    validators,
    col_widths=[2.8, 3.7],
)

body("ValidationResult dataclass:", bold=True)
code_block("{ patient_ids: list[str], drug_name: str | None, dose: str | None }")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. GUARDRAILS — DEFENSE IN DEPTH
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Guardrails — Defense in Depth (7 Layers)")

h2("Layer 1 — Prompt Injection Defense")
body("File: src/chatbot/intent_classifier.py — is_prompt_injection(text)", bold=True)
body("Runs before any LLM call. Uses regex to detect adversarial patterns:")
for p in ["ignore rbac / bypass rbac / disable rbac",
          "show all patient / dump all / list all records",
          "system prompt / jailbreak / ignore previous instructions"]:
    bullet(p)
body("Result: intent tagged rbac_sensitive_request → denied at controller level with zero LLM invocation.")

h2("Layer 2 — PHI Request Denial")
body("File: src/chatbot/phi_guard.py — deny_if_phi_requested(user_text)", bold=True)
body("Blocked phrases:")
for p in ["full patient profile", "patient name", "date of birth", "dob", "mrn",
          "medical record", "discharge note", "doctor details", "attending physician"]:
    bullet(p)
body("Blocked regex patterns: NAME OF PAT-\\d{3} or PAT-\\d{3}.*NAME (unless drug-related context).")
body("Action: raises PHIError(PermissionError) immediately. No LLM invoked.")

h2("Layer 3 — PHI Field Detection (Recursive)")
body("File: src/chatbot/phi_guard.py — contains_phi_keys(data)", bold=True)
body("Walks any nested dict/list structure with full path tracking:")
make_table(
    ["Field Category", "Fields"],
    [
        ["Always-PHI", "dob, mrn, discharge_note, attending_physician"],
        ["Contextual-PHI", "name (only PHI when in patient context — path contains 'patient'/'demographics' or container has patient_id/mrn/dob keys)"],
    ],
    col_widths=[2.0, 4.5],
)

h2("Layer 4 — PHI Stripping")
body("strip_phi(data) — recursive removal returning a clean dict. Applied by:", bold=True)
bullet("EHRServer.get_billing_safe_summary() — strips before returning to billing")
bullet("Telemetry._scrub_phi() — strips before any observability write")
body("Fields removed: name, dob, mrn, discharge_note, attending_physician")

h2("Layer 5 — Billing-Side PHI Validation")
body("File: src/servers/billing_server.py — validate_billing_data(payload)", bold=True)
body("Validates incoming data has NO PHI fields. If PHI found → raises PHIBoundaryViolationError(field, server, payload_sample). Called by generate_invoice() before processing any billing data.")

h2("Layer 6 — RBAC Enforcement (Per Tool Call)")
body("File: src/utils/rbac.py — RBACEngine.check_permission(role, server, tool, patient_id)", bold=True)
code_block("""\
Every MCP tool call:
  rbac_guard.RBACGuard.ensure_allowed(actor, server, tool, patient_id)
    → rbac_engine.check_permission(role, server, tool, patient_id)
      → TOOL_PERMISSIONS[tool] → required_permission
      → if required_permission not in policy[role][server]:
            telemetry.record_rbac_violation(...)
            raise RBACError(role, tool, server)""")

h2("Layer 7 — Input Sanitization")
body("File: src/chatbot/preprocessor.py — sanitize_user_text()", bold=True)
for item in [
    "Unicode normalisation (NFKC)",
    "Zero-width character removal (\\u200b, \\u200c, etc.)",
    "Control character replacement",
    "Smart quote → ASCII (“” → \", ‘’ → ')",
    "Whitespace collapse",
    "Patient ID normalisation: pat001 / pat 1 / PAT-1 → PAT-001",
    "Typo corrections: dishcharge→discharge, medecine→medicine",
]:
    bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. MONITORING & TELEMETRY
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Monitoring & Telemetry")

h2("Telemetry Singleton — src/utils/telemetry.py")
body("Centralized observability store (in-memory; production: migrate to Redis).")
body("All data written through _scrub_phi() — zero PHI in any telemetry record.", bold=True)

h3("MCPCall dataclass (every tool invocation)")
code_block("timestamp, server, tool, role, patient_id, duration_ms, success: bool, error: str|None")

h3("ChatTrace dataclass (every user request)")
code_block("""\
timestamp, conversation_id, role, patient_id,
latency_ms, success: bool, mcp_calls: int,
rbac_violations: int, needs_clarification: bool,
clarification_type: str|None, error: str|None""")

h3("Alert dataclass")
code_block('timestamp, level: "INFO"|"WARNING"|"ERROR"|"CRITICAL"|"URGENT", source, message, details')

h3("MCPCallTimer Context Manager")
body("Every tool invocation is automatically wrapped:")
code_block("""\
with MCPCallTimer("ehr", "get_discharge_medications", role, patient_id):
    return server_impl.get_discharge_medications(patient_id, role)

# Captures: wall-clock duration, success/failure, per-server + per-tool counts""")

h3("get_summary() response shape")
code_block("""\
{
  "total_calls": 42,
  "uptime_seconds": 3600,
  "by_server": {
    "ehr":      {"success": 15, "failure": 0},
    "pharmacy": {"success": 20, "failure": 1},
    "billing":  {"success":  7, "failure": 0}
  },
  "alerts_by_level": {"INFO": 3, "WARNING": 1, "ERROR": 0, "CRITICAL": 0},
  "rbac_violations": 2
}""")

h2("LangSmith Tracing — src/utils/langsmith_tracing.py (411 lines)")
body("Enabled by LANGCHAIN_TRACING_V2=true + LANGSMITH_API_KEY. Safe-by-default: no-op if unconfigured.")

make_table(
    ["Processor", "What It Sanitizes"],
    [
        ["process_inputs_controller",   "Masks user text, preserves patient IDs"],
        ["process_outputs_controller",  "Summarises answer, masks text"],
        ["process_inputs_llm_provider", "Logs message/tool count only (no content)"],
        ["process_outputs_llm_provider","Logs response shape, checks for PHI"],
        ["process_inputs_mcp_retry",    "Sanitizes MCP call args"],
        ["process_outputs_mcp_retry",   "Sanitizes MCP results"],
        ["process_inputs_workflow",     "Sanitizes workflow inputs"],
        ["process_outputs_workflow",    "Sanitizes workflow outputs"],
        ["process_inputs_mcp_tool",     "Sanitizes FastMCP tool args"],
        ["process_outputs_mcp_tool",    "Sanitizes FastMCP results"],
    ],
    col_widths=[2.8, 3.7],
)

body("instrument_fastmcp_tools(mcp, server) wraps all subsequently-registered FastMCP tools with LangSmith tracing automatically.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. EVALUATION SUITE
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Evaluation Suite")

make_table(
    ["File", "What It Tests"],
    [
        ["eval_dashboard.py",                              "Runs full discharge for all 6 patients; collects latency, tool call counts, alerts, substitution rates, success rate"],
        ["test_chatbot_flows.py",                         "Multi-turn conversations: follow-up queries, context retention, ordinal references"],
        ["test_llm_agent_guards.py",                      "PHI request attempts + RBAC violation attempts — verifies they are blocked correctly"],
        ["test_phi_guard_invoice_payloads.py",            "Sends invoices with PHI fields included — verifies PHIBoundaryViolationError is raised"],
        ["test_llm_agent_tool_name_normalization.py",     "Drug alias resolution (brand → generic, alias → drug ID)"],
        ["test_patient_ordinal_resolution.py",            "'First patient', '3rd patient', 'second' → correct PAT-XXX mapping"],
    ],
    col_widths=[3.0, 3.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. SERVER GATEWAY
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Server Gateway")

h2("FastAPI Chat Gateway — src/gateway/chat_gateway.py (port 8000)")

make_table(
    ["Method", "Route", "Purpose"],
    [
        ["POST", "/chat",                    "Main chat: user text → controller → formatted response"],
        ["GET",  "/sessions",                "List all active conversation sessions"],
        ["GET",  "/sessions/{conv_id}",      "Get full history of a session"],
        ["POST", "/sessions/{conv_id}/clear","Clear conversation context"],
        ["POST", "/invoice/{patient_id}",    "Generate + return invoice PDF (binary download)"],
        ["POST", "/prescription/{patient_id}","Generate + return prescription PDF"],
        ["GET",  "/invoice/{patient_id}/html","HTML preview of invoice"],
        ["GET",  "/health",                  "System health: server TCP pings + uptime"],
    ],
    col_widths=[0.7, 2.3, 3.5],
)

h3("Session Management")
body("SESSION_STORE: OrderedDict capped at 50. Each session stores: conversation_id, exchanges list [{user, assistant, data, timestamp}].")

h3("Invoice Collection Pipeline")
code_block("""\
_collect_invoice_data(patient_id)
  → MCPClient.ehr_call("get_billing_safe_summary")
  → MCPClient.billing_call("get_charges")
  → MCPClient.billing_call("get_insurance")
  → MCPClient.pharmacy_call("get_bulk_price")
  → returns {billing_safe, charges, drug_prices}""")

h2("MCP Client Proxy — src/chatbot/mcp_client.py")
for item in [
    "Lazy connection on first tool call",
    "Per-server asyncio.Lock() (SSE not concurrency-safe)",
    "RBAC enforcement before every call",
    "Exponential backoff retry (max 3 attempts + jitter)",
    "Telemetry recording: duration, success, role, patient_id",
    "Wraps failures in ToolExecutionError",
]:
    bullet(item)

h2("LLM Provider Gateway — src/gateway/llm_azure.py")
for item in [
    "AzureOpenAI client initialisation from env vars",
    "chat_with_tools(messages, tools, temperature) — async",
    "Supports parallel tool calls (fallback to sequential if unsupported)",
    "LangSmith-traced with safe input/output processors",
]:
    bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. MODEL CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Model Configuration")

h2("Primary LLM — Azure OpenAI")
make_table(
    ["Setting", "Value"],
    [
        ["Endpoint",     "https://akshata-openai-keys.openai.azure.com/"],
        ["Deployment",   "gpt-5.4"],
        ["API Version",  "2024-12-01-preview"],
        ["Temperature",  "0.2 (deterministic tool calling)"],
        ["Tool Choice",  "auto"],
        ["Parallel Calls","True (fallback to sequential if unsupported)"],
    ],
    col_widths=[2.5, 4.0],
)

h2("Secondary LLM — Preprocessor (Azure OpenAI)")
make_table(
    ["Setting", "Value"],
    [
        ["Endpoint",   "https://clinical-decision-system.openai.azure.com/"],
        ["Deployment", "gpt-4o"],
        ["Temperature","0.0 (fully deterministic)"],
        ["Trigger",    "CHATBOT_USE_LLM_PREPROCESSOR=1 env var"],
        ["Note",       "Typo in .env: ZURE_OPENAI_API_VERSION_ONE (missing leading A) may prevent init"],
    ],
    col_widths=[2.5, 4.0],
)

h2("Tool Calling Flow")
code_block("""\
llm_provider.chat_with_tools(messages, tools, temperature)
  → AzureOpenAI(endpoint, api_key, api_version)
  → client.chat.completions.create(
        model=deployment_name,
        messages=messages,
        tools=tools,
        tool_choice="auto",
        parallel_tool_calls=True
    )
  → returns tool_calls list or content string""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. RBAC SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
h1("12. RBAC System")

h2("RBACEngine Singleton — src/utils/rbac.py")
body("Policy source: data/rbac_policies.json. Zero code changes to add/modify roles.")

h2("Role Access Matrix")
make_table(
    ["Role", "EHR", "Pharmacy", "Billing", "Security", "Telemetry"],
    [
        ["discharge_coordinator", "Full read", "Full read", "Invoice + charges", "Read", "Read"],
        ["billing_agent",         "Diagnosis + admission dates only", "Price only", "Full billing", "Read", "Read"],
        ["pharmacy_agent",        "Medications + diagnoses", "Full", "BLOCKED", "Read", "Read"],
        ["clinical_agent",        "Full clinical", "Medications", "BLOCKED", "Read", "Read"],
    ],
    col_widths=[1.6, 1.3, 1.3, 1.3, 1.0, 1.0],
)

h2("Tool → Permission Mapping (sample from 50+ entries)")
code_block("""\
"get_discharge_medications"    → "read_medications"
"get_billing_safe_summary"     → "read_billing_safe"
"get_diagnosis_codes"          → "read_diagnosis_codes"
"get_admission_info"           → "read_admission_dates"
"check_stock"                  → "check_stock"
"get_price"                    → "get_drug_price"
"generate_invoice"             → "generate_invoice"
"validate_billing_data"        → "validate_billing"
"get_mcp_call_count"           → "read_telemetry"
... and 40+ more mappings""")

h2("Violation Log Entry Schema")
code_block("""\
{
  "timestamp":  "2024-01-15T10:23:41.123Z",
  "role":       "billing_agent",
  "server":     "pharmacy",
  "tool":       "update_stock",
  "patient_id": "PAT-001"
}""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. PHI GUARD
# ══════════════════════════════════════════════════════════════════════════════
h1("13. PHI Guard")

h2("src/chatbot/phi_guard.py")

h3("PHI Field Classification")
make_table(
    ["Category", "Fields"],
    [
        ["ALWAYS_PHI_FIELDS",     "dob, mrn, discharge_note, attending_physician"],
        ["CONTEXTUAL_PHI_FIELDS", "name (only PHI when in patient context)"],
        ["PHI_FIELDS",            "Union of both sets"],
    ],
    col_widths=[2.5, 4.0],
)

h3("Three Defense Layers")

body("Layer A — Request Denial (pre-LLM):", bold=True)
code_block("""\
deny_if_phi_requested("what is the name of PAT-001?")
  keyword scan: "patient name", "full profile", "dob", "mrn", …
  regex: NAME OF PAT-\\d{3} or PAT-\\d{3}.*NAME
  → raises PHIError("Request denied. PHI cannot be disclosed.")""")

body("Layer B — Recursive Payload Scan:", bold=True)
code_block("""\
contains_phi_keys({"patient": {"name": "John", "ward": "ICU"}})
  walks all keys recursively with path tracking
  _is_person_context(["patient","name"], container)
      path has "patient" token → YES → "name" is PHI here
  returns True (PHI found)""")

body("Layer C — PHI Stripping:", bold=True)
code_block("""\
strip_phi({"patient_id":"PAT-001","name":"John","ward":"ICU","dob":"1990-01-01"})
  returns {"patient_id":"PAT-001","ward":"ICU"}
  name + dob removed""")

h3("EHR-Side PHI Boundary (get_billing_safe_summary)")
body("Fields removed before returning to billing layer:")
for f in ["name", "dob", "mrn", "discharge_note", "attending_physician"]:
    bullet(f)
body("Fields retained: patient_id, ward, admission_date, discharge_date, los_days, diagnosis_icd10, special_instructions")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. RESPONSE FORMATTER
# ══════════════════════════════════════════════════════════════════════════════
h1("14. Response Formatter")
h2("src/chatbot/response_formatter.py — 361 lines")

make_table(
    ["Function", "Output Style"],
    [
        ["format_medication_list(patient_id, meds)",             "Prescribed medications for PAT-001:\\n1. Aspirin 100mg once daily…"],
        ["format_medication_lists(meds_by_patient)",             "Multi-patient grouped list"],
        ["format_all_patients_report(sections_by_patient)",      "Full report spanning all 6 patients"],
        ["format_stock_check_list(patient_id, result, summary)", "✔ Available / ⚠ Not Available + alternatives + dose conflicts + doctor disclaimer"],
        ["format_unavailable_only(patient_id, result)",          "Only out-of-stock drugs + alternatives"],
        ["format_success_discharge(result)",                     "Success summary with alerts + substitutions + invoice total"],
        ["format_discharge_summary_safe(patient_id, **kwargs)",  "8-section PHI-safe discharge summary: Admission, Diagnosis, Medications, Substitutions, Clinical Review, Alerts, Stock, Invoice"],
        ["format_access_denied()",                               "'Access denied. Not authorised to view full clinical discharge summaries.'"],
        ["format_phi_denied()",                                  "'Request denied. Sensitive patient information (PHI) cannot be included in billing.'"],
        ["format_observability(patient_id, summary)",            "'Total calls: X, Alerts: Y, RBAC violations: Z'"],
    ],
    col_widths=[2.8, 3.7],
)

h3("Sample Stock Check Output")
code_block("""\
Stock Check Result for PAT-003:

✔ Available:
  • Aspirin 100mg → 50 units in stock

⚠ Not Available:
  • Furosemide 40mg

Suggested Alternatives:
  • Furosemide 40mg → Torsemide 10mg (available, same class)

⚠ Please consult your doctor before switching to an alternative medication.""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 15. METRICS SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
h1("15. Metrics System")

h2("RequestMetrics — src/chatbot/metrics.py")
code_block("""\
@dataclass
class RequestMetrics:
    patient_id: str
    mcp_call_count: int = 0
    mcp_call_count_by_server: dict = field(default_factory=dict)
    alerts: list = field(default_factory=list)
    rbac_violations: int = 0

    def add_alert(self, alert: str): ...""")

h2("DischargeCoordinationAgent Metrics — src/agents/discharge_agent.py")
body("orchestrate_discharge() returns dict with metrics keys:", bold=True)
code_block("""\
{
  "total_latency_ms":          823,
  "mcp_tool_calls_total":      11,
  "mcp_tool_calls_success":    10,
  "mcp_tool_calls_failed":     1,
  "tool_call_success_rate":    0.909,
  "phi_boundary_enforced":     true,
  "phi_blocked_fields":        ["name","dob","mrn"],
  "alerts":                    ["DOSE_CONFLICT: Metformin 500mg vs 1000mg standard"],
  "data_integrity_conflicts":  []
}""")

h2("WorkflowMetrics — Benchmark Constants")
make_table(
    ["Metric", "Manual Process", "MCP Automated"],
    [
        ["Handoffs",       "15",         "0 (automated)"],
        ["Average time",   "45 minutes", "<1 second"],
        ["Error rate",     "18%",        "~0% (guardrails)"],
    ],
    col_widths=[2.0, 2.0, 2.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 16. CUSTOM EXCEPTION HIERARCHY
# ══════════════════════════════════════════════════════════════════════════════
h1("16. Custom Exception Hierarchy")
h2("src/utils/exceptions.py")
code_block("""\
Exception
└── MCPDischargeError(message, details)
    ├── RBACError(role, tool, server)
    │     "Role 'billing_agent' cannot call 'update_stock' on 'pharmacy'"
    ├── PHIBoundaryViolationError(field, server, payload_sample)
    │     "PHI field 'name' found in billing payload"
    ├── StockUnavailableError(drug_name, patient_id, alternatives)
    │     "Furosemide 40mg out of stock for PAT-003"
    ├── DoseConflictError(drug1, drug2, reason)
    │     "Metformin 500mg vs formulary standard 1000mg"
    ├── MCPConnectionError(server, reason)
    │     "Cannot reach EHR server at localhost:8001"
    └── ToolExecutionError(tool, server, reason)
          "get_discharge_medications failed: patient not found\"""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 17. DATA LAYER
# ══════════════════════════════════════════════════════════════════════════════
h1("17. Data Layer")

h2("JSON Data Files (All 7)")
make_table(
    ["File", "Size", "Contents"],
    [
        ["ehr_patients.json",         "9.8 KB",  "6 patients — patient_id, mrn, name, dob, ward, admission_date, discharge_date, los_days, attending_physician, diagnosis_icd10, discharge_note, discharge_medications, special_instructions"],
        ["pharmacy_inventory.json",   "13 KB",   "17 drugs — drug_id, generic_name, brand_names, therapeutic_class, strengths, formulations, in_stock, stock_units, reorder_threshold, price_per_unit_inr, formulary_standard_dose, aliases"],
        ["billing_rate_cards.json",   "2.1 KB",  "15 charge codes — charge_code, name, rate_inr, ward_type, applies_to_los_days"],
        ["insurance_contracts.json",  "922 B",   "2 insurers — insurer_id, plan_type, copay_inr, deductible_inr, max_covered_per_admission_inr, covered_icd10_prefixes"],
        ["patient_insurance_map.json","689 B",   "Patient → insurer mappings with policy_number + pa_required"],
        ["icd10_billing_codes.json",  "984 B",   "ICD-10 → DRG billing code mappings"],
        ["rbac_policies.json",        "2.1 KB",  "RBAC matrix: {role: {server: [permissions]}} for 4 roles × 5 servers"],
    ],
    col_widths=[2.2, 0.7, 3.6],
)

h3("Out-of-Stock Drugs (4)")
for d in ["Furosemide 40mg", "Adalimumab / Humira", "Tafamidis / Vyndamax", "Osimertinib / Tagrisso"]:
    bullet(d)

body("These trigger the get_alternative() tool call path and generate STOCK_UNAVAILABLE alerts.", italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 18. SEMANTIC DRUG MATCHING
# ══════════════════════════════════════════════════════════════════════════════
h1("18. Semantic Drug Matching")
h2("src/servers/pharmacy_server.py — _resolve_drug()")
body("Algorithm: word-overlap scoring (no embeddings; deterministic).")
code_block("""\
_resolve_drug("Farxiga 10mg")
  1. Normalize: lowercase, strip dose
  2. Check aliases table: "farxiga" → "PH-001" (Dapagliflozin)
  3. Score word overlap against all drug names + aliases
  4. Require ≥ 85% overlap to avoid "NAME_MISMATCH" alert
  5. Return drug record or None""")

body("Alias table in pharmacy_inventory.json supports: generic name, all brand names, abbreviations, common misspellings.")
body("NAME_MISMATCH alert is generated when overlap < 85% — flags for clinical review.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 19. KEY DESIGN PATTERNS
# ══════════════════════════════════════════════════════════════════════════════
h1("19. Key Design Patterns")
make_table(
    ["Pattern", "Implementation", "Location"],
    [
        ["Singleton",                "RBACEngine, Telemetry, DataLoader — get_*() factory functions", "src/utils/"],
        ["Context Manager",          "MCPCallTimer, AsyncMCPToolClient — automatic resource management", "telemetry.py, mcp_client.py"],
        ["Exponential Backoff",      "_retry_call(client, tool, args) max 3 attempts + random jitter", "mcp_client.py, discharge_agent.py"],
        ["PHI Boundary Enforcement", "Strip at EHR source (get_billing_safe_summary) + validate at Billing sink (validate_billing_data)", "ehr_server.py, billing_server.py"],
        ["Defense in Depth",         "7 independent guardrail layers — each independently blocks attacks", "llm_controller.py"],
        ["Data-Driven RBAC",         "rbac_policies.json — zero code changes to add/modify roles or permissions", "rbac.py"],
        ["SSE Concurrency Safety",   "Per-server asyncio.Lock() — prevents race conditions on SSE connections", "mcp_client.py"],
        ["Safe-by-Default Tracing",  "LangSmith no-op if unconfigured; graceful degradation on errors", "langsmith_tracing.py"],
        ["Tool-Level RBAC",          "Every MCP tool call validates permission before any data access", "All servers"],
        ["Fallback Routing",         "LLM timeout → rule-based routing (_fallback) ensures availability", "llm_controller.py"],
    ],
    col_widths=[1.8, 2.7, 2.0],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 20. CONFIGURATION & ENVIRONMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("20. Configuration & Environment Variables")
make_table(
    ["Variable", "Purpose"],
    [
        ["AZURE_OPENAI_ENDPOINT",           "Primary LLM endpoint URL"],
        ["AZURE_OPENAI_API_KEY",            "Primary LLM API key"],
        ["AZURE_OPENAI_DEPLOYMENT_NAME",    "Primary deployment (gpt-5.4)"],
        ["AZURE_OPENAI_API_VERSION",        "API version (2024-12-01-preview)"],
        ["AZURE_OPENAI_MODEL_NAME",         "Model name for client initialisation"],
        ["AZURE_OPENAI_ENDPOINT_ONE",       "Secondary LLM endpoint (preprocessor)"],
        ["AZURE_OPENAI_API_KEY_ONE",        "Secondary LLM API key"],
        ["AZURE_OPENAI_DEPLOYMENT_NAME_ONE","Secondary deployment (gpt-4o)"],
        ["ZURE_OPENAI_API_VERSION_ONE",     "⚠ Typo — missing leading A. Secondary API version may not load"],
        ["LANGCHAIN_TRACING_V2",            "Enable LangSmith tracing (true/false)"],
        ["LANGCHAIN_API_KEY",               "LangSmith API key"],
        ["LANGCHAIN_PROJECT",               "LangSmith project name (MCPDischarge)"],
        ["CHATBOT_USE_LLM_PREPROCESSOR",    "Set to 1 to enable LLM text rewriting"],
    ],
    col_widths=[2.8, 3.7],
)

h2("Python Dependencies (requirements.txt)")
make_table(
    ["Package", "Purpose"],
    [
        ["fastmcp>=0.1.0",                  "FastMCP HTTP SSE server framework"],
        ["mcp>=1.0.0",                      "Official MCP Python SDK (client + server)"],
        ["fastapi>=0.110.0",                "HTTP gateway framework"],
        ["uvicorn[standard]>=0.27.0",       "ASGI server for FastAPI"],
        ["openai>=1.30.0",                  "Azure OpenAI SDK"],
        ["langchain>=0.2.0",               "LangChain orchestration"],
        ["langchain-openai>=0.1.0",         "LangChain Azure OpenAI integration"],
        ["langsmith>=0.7.0",               "LangSmith tracing SDK"],
        ["azure-ai-projects>=1.0.0b1",      "Azure AI Foundry integration"],
        ["azure-identity>=1.17.0",          "Azure credential management"],
        ["azure-monitor-opentelemetry>=1.0","Azure monitoring + OpenTelemetry"],
        ["reportlab>=4.1.0",               "PDF generation (invoices, prescriptions)"],
        ["python-dotenv>=1.0.0",           "Environment variable loading"],
        ["jinja2>=3.1.3",                  "HTML template rendering"],
        ["rich>=13.7.0",                   "Terminal formatting for CLI"],
        ["numpy, pandas, matplotlib",      "Data processing + evaluation charts"],
        ["jsonschema>=4.20.0",             "JSON schema validation"],
    ],
    col_widths=[2.5, 4.0],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 21. NUMBERS AT A GLANCE
# ══════════════════════════════════════════════════════════════════════════════
h1("21. Numbers at a Glance")
make_table(
    ["Category", "Count / Value"],
    [
        ["MCP Tools (total)",           "59"],
        ["EHR tools",                   "18"],
        ["Pharmacy tools",              "18"],
        ["Billing tools",               "11"],
        ["Security tools",              "4"],
        ["Telemetry tools",             "8"],
        ["Python source files",         "25+"],
        ["Lines of code",               "~3,500"],
        ["Custom exception types",      "7"],
        ["RBAC roles",                  "4"],
        ["RBAC permission types",       "30+"],
        ["PHI fields guarded",          "5 (name, dob, mrn, discharge_note, attending_physician)"],
        ["Guardrail layers",            "7"],
        ["Test patients",               "6 (PAT-001 – PAT-006)"],
        ["Drug records",                "17 (4 out-of-stock)"],
        ["Billing charge codes",        "15"],
        ["Insurance companies",         "2"],
        ["Evaluation test files",       "6"],
        ["MCP servers",                 "5 (ports 8001–8005)"],
        ["Gateway port",                "8000"],
        ["LangSmith processor funcs",   "10"],
        ["Manual process time (before)","45 minutes, 15 handoffs, 18% error rate"],
        ["Automated time (after)",      "<1 second, 0 handoffs, ~0% error rate"],
    ],
    col_widths=[3.0, 3.5],
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\pramo\Downloads\project5_mcp\MCPDischarge_Project_Analysis.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
